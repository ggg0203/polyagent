"""生成 PolyAgent 产品说明书 (.docx)."""

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn

doc = Document()

# ── 全局样式 ──────────────────────────────────────────
style = doc.styles["Normal"]
font = style.font
font.name = "Microsoft YaHei"
font.size = Pt(11)
style.element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")

for level in range(1, 4):
    hs = doc.styles[f"Heading {level}"]
    hs.font.name = "Microsoft YaHei"
    hs.element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    hs.font.color.rgb = RGBColor(0x1A, 0x56, 0xDB)

# ── 封面 ──────────────────────────────────────────────
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("\n\n\n\n")
run.font.size = Pt(36)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("PolyAgent")
run.bold = True
run.font.size = Pt(36)
run.font.color.rgb = RGBColor(0x1A, 0x56, 0xDB)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("模型无关的多智能体协作框架")
run.font.size = Pt(18)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("\n\n产品说明书\n")
run.font.size = Pt(14)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("版本 0.1.0 — 2026 年 7 月")
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

doc.add_page_break()

# ── 目录页 ────────────────────────────────────────────
doc.add_heading("目录", level=1)
toc_items = [
    "1. 产品概述",
    "2. 核心架构",
    "3. 快速开始",
    "4. CLI 使用指南",
    "5. SDK 使用指南",
    "6. 内置工具一览",
    "7. 技能市场",
    "8. 多 Agent 编排",
    "9. 企业级特性",
    "10. 测试与质量保证",
    "11. 常见问题",
    "12. 附录",
]
for item in toc_items:
    p = doc.add_paragraph(item, style="List Number")
    p.paragraph_format.space_after = Pt(4)

doc.add_page_break()

# ══════════════════════════════════════════════════════
# 1. 产品概述
# ══════════════════════════════════════════════════════
doc.add_heading("1. 产品概述", level=1)

doc.add_heading("1.1 什么是 PolyAgent", level=2)
doc.add_paragraph(
    "PolyAgent 是一个纯后端、模型无关的多智能体协作框架。它不依赖任何特定的 LLM 提供商，"
    "核心定位是：将多个智能体（Agent）组织成一条可编排、可观测、可测评的生产流水线。"
    "PolyAgent 提供 SDK（Python 库）和 CLI（命令行工具）两种交付形态，适合嵌入到任何"
    "需要多智能体协作能力的后端系统中。"
)

doc.add_heading("1.2 设计动机", level=2)
doc.add_paragraph(
    "在实际工作中，单 Agent 直接调用 LLM 会碰到三个瓶颈："
)
problems = [
    ("上下文爆炸", "一个 Agent 同时做计划 + 执行 + 总结，prompt 不断膨胀"),
    ("单点不可靠", "一次超时 / 限流就搞砸整个任务"),
    ("不可观测", "无法知道花了多少 Token、花了多少钱、效果如何"),
]
for title, desc in problems:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    run = p.add_run(f"• {title}：")
    run.bold = True
    p.add_run(desc)

doc.add_paragraph(
    "PolyAgent 通过角色分工 + 编排层 + 可靠性中间件链来解决这些问题，"
    "把原始的 LLM 调用变成了一个可调度、可降解、可观测的系统。"
)

doc.add_heading("1.3 关键特性", level=2)
features = [
    "多 Agent 编排：Planner → Workers(并行) → Critic(评审回退) → Synthesizer",
    "模型无关：支持 Mock（离线测试）、DeepSeek、以及任何 OpenAI 兼容协议",
    "工具系统：代码执行、文件读写、HTTP 请求、文件搜索、联网搜索",
    "双重技能市场：自建仓库 + 腾讯 SkillHub（76,000+ 技能）",
    "可观测：分布式 Tracing（Jaeger/OTLP）、Metrics（Prometheus）",
    "内存系统：Buffer 记忆、向量记忆、压缩",
    "RAG：FastEmbed（ONNX 本地语义）+ ChromaDB 持久化",
    "Docker 沙箱：隔离执行不可信代码",
    "评估框架：自动打分（ContainsScorer / LLM Judge）",
    "质量门禁：mypy strict、ruff、覆盖率 75%",
]
for f in features:
    doc.add_paragraph(f, style="List Bullet")

doc.add_page_break()

# ══════════════════════════════════════════════════════
# 2. 核心架构
# ══════════════════════════════════════════════════════
doc.add_heading("2. 核心架构", level=1)

doc.add_heading("2.1 分层架构", level=2)
doc.add_paragraph(
    "PolyAgent 采用清晰的分层架构，从上到下依次为："
)

layers = [
    ("CLI 层", "polyagent run / chat / shell / talk / eval / runs / show 命令入口"),
    ("编排层", "Orchestrator 核心：Planner（计划）→ Workers（并行执行）→ Critic（评审回退）→ Synthesizer（汇总）"),
    ("Agent 执行器", "单 Agent 的工具循环 + 中间件链（限流→重试→降级→预算校验→成本统计）"),
    ("LLM 层", "MockProvider（离线）/ DeepSeekProvider（真实），支持流式输出"),
    ("工具系统", "PythonExecute / ReadFile / WriteFile / HttpRequest / GrepFiles / WebSearch"),
    ("记忆 + RAG", "BufferMemory / VectorMemory / Compressor + FastEmbed + ChromaDB"),
    ("可观测", "Tracer（Span 树）/ Metrics / OTLP 导出 / Prometheus 推送"),
    ("持久化", "SQLite RunStore（保存执行历史）/ .env 配置"),
]
for title, desc in layers:
    p = doc.add_paragraph()
    run = p.add_run(f"{title}：")
    run.bold = True
    p.add_run(desc)

doc.add_heading("2.2 编排流水线", level=2)
doc.add_paragraph(
    "编排流水线是 PolyAgent 的核心价值。当用户运行 polyagent run 或通过 SDK 调用 Orchestrator 时，"
    "系统执行以下步骤："
)

pipeline = [
    "Planner：将用户目标分解为 DAG（有向无环图）子任务列表",
    "Workers：根据依赖拓扑并发执行子任务，可带工具",
    "Critic：评审每个子任务的输出质量，不通过则回退重做（max_review_retries）",
    "Synthesizer：汇总所有子任务结果，生成最终答案",
]
for i, step in enumerate(pipeline, 1):
    doc.add_paragraph(f"{i}. {step}")

doc.add_page_break()

# ══════════════════════════════════════════════════════
# 3. 快速开始
# ══════════════════════════════════════════════════════
doc.add_heading("3. 快速开始", level=1)

doc.add_heading("3.1 环境要求", level=2)
doc.add_paragraph("• Python ≥ 3.11")
doc.add_paragraph("• pip install -e .[dev]（可编辑安装 + 开发工具）")
doc.add_paragraph("• 可选：pip install -e .[rag]（RAG 扩展）")
doc.add_paragraph("• 可选：配置 .env 文件中的 DEEPSEEK_API_KEY")

doc.add_heading("3.2 运行前准备", level=2)

p = doc.add_paragraph()
run = p.add_run("⚠️ 重要：所有 CLI 命令必须在项目根目录下运行。")
run.bold = True
run.font.color.rgb = RGBColor(0xCC, 0x33, 0x33)

doc.add_paragraph(
    "PolyAgent 安装在 venv 的 Scripts 目录中，不在系统 PATH 里。"
    "同时，框架需要在项目根目录读取 .env 配置文件。"
    "所以首次使用前需要两步准备："
)

doc.add_paragraph("第一步：打开 CMD 并切换到项目根目录")
p = doc.add_paragraph('  cd /d D:\\孙\\WorkBuddy\\2026-07-10-18-40-02')
for r in p.runs:
    r.font.name = "Consolas"
    r.font.size = Pt(10)

doc.add_paragraph("第二步：将 venv 的 Scripts 目录临时加入 PATH")
p = doc.add_paragraph(
    '  set PATH=C:\\Users\\孙\\.workbuddy\\binaries\\python\\envs\\default\\Scripts;%PATH%'
)
for r in p.runs:
    r.font.name = "Consolas"
    r.font.size = Pt(10)

doc.add_paragraph(
    "完成这两步后，就可以直接使用 polyagent 命令了。"
    "你可以在 CMD 中输入 polyagent --help 来查看所有可用命令。"
)

doc.add_paragraph(
    '如果你想永久将 polyagent 加入 PATH，可以：'
    '右键「此电脑」→ 属性 → 高级系统设置 → 环境变量 → '
    '在「系统变量」的 Path 中新增：'
    'C:\\Users\\孙\\.workbuddy\\binaries\\python\\envs\\default\\Scripts。'
    '设置后重新打开 CMD 即可全局使用。'
)

doc.add_heading("3.3 离线体验", level=2)
doc.add_paragraph(
    "无需任何 API Key，直接 CLI 体验编排流水线："
).runs[0].font.name = "Consolas"

codes = [
    "polyagent version              # 查看版本",
    'polyagent run "写一个简单应用"   # 运行多 Agent 流水线（Mock）',
    "polyagent eval                 # 跑评测数据集",
    "polyagent chat                 # 交互式单 Agent 对话（Mock）",
    "polyagent talk                 # 聊天模式（拟人化，无需 Key）",
]
for c in codes:
    p = doc.add_paragraph(c)
    for run in p.runs:
        run.font.name = "Consolas"
        run.font.size = Pt(10)

doc.add_heading("3.4 联网体验", level=2)
doc.add_paragraph("配置好 .env 后，使用 DeepSeek 真实 LLM：")

codes2 = [
    "set PATH=C:\\Users\\孙\\.workbuddy\\binaries\\python\\envs\\default\\Scripts;%PATH%",
    "cd /d D:\\孙\\WorkBuddy\\2026-07-10-18-40-02    # 切换到项目根目录",
    "polyagent shell                  # 启动带工具的智能助手",
    "polyagent talk                   # 启动聊天模式（拟人化，隐藏工具）",
    "polyagent run \"分析这个仓库\" --provider deepseek",
    "polyagent chat --provider deepseek --stream",
]
for c in codes2:
    p = doc.add_paragraph(c)
    for run in p.runs:
        run.font.name = "Consolas"
        run.font.size = Pt(10)

doc.add_page_break()

# ══════════════════════════════════════════════════════
# 4. CLI 使用指南
# ══════════════════════════════════════════════════════
doc.add_heading("4. CLI 使用指南", level=1)

commands = [
    ("version", "查看版本号", "polyagent version"),
    ("run", "运行多 Agent 编排流水线", 'polyagent run "目标" [--provider deepseek] [--persist/--no-persist]'),
    ("chat", "交互式单 Agent 对话", "polyagent chat [--provider deepseek] [--stream]"),
    ("shell", "带全套工具的智能助手（推荐）", "polyagent shell"),
    ("talk", "聊天模式—拟人化对话，隐藏工具调用过程", "polyagent talk"),
    ("eval", "运行评测数据集", "polyagent eval"),
    ("runs", "列出历史运行记录", "polyagent runs [--limit 20]"),
    ("show", "查看某次运行的详细信息", "polyagent show <run_id>"),
]

table = doc.add_table(rows=1, cols=3)
table.style = "Light Grid Accent 1"
table.alignment = WD_TABLE_ALIGNMENT.CENTER

hdr = table.rows[0].cells
hdr[0].text = "命令"
hdr[1].text = "说明"
hdr[2].text = "示例"

for cmd, desc, example in commands:
    row = table.add_row().cells
    row[0].text = cmd
    row[1].text = desc
    row[2].text = example

doc.add_paragraph()  # spacing

doc.add_heading("4.1 shell 命令详解", level=2)
doc.add_paragraph(
    "shell 是推荐使用的交互模式。它给 Agent 挂载了全部 12 个内置工具："
    "python_execute / read_file / write_file / http_request / grep_files / "
    "web_search / marketplace_search / marketplace_install / "
    "skillhub_search / skillhub_install / skillhub_install_from_prompt / skillhub_list。"
    "Agent 会根据你的需求智能选择工具，并实时显示调用过程。"
)

doc.add_heading("4.2 run 命令详解", level=2)
doc.add_paragraph(
    "run 命令执行完整的多人编排流水线（Planner → Workers → Critic → Synthesizer），"
    "适用于一次性任务。支持 --persist 参数自动保存执行历史到 SQLite 数据库。"
)

doc.add_heading("4.3 环境配置", level=2)
doc.add_paragraph(
    "项目根目录的 .env 文件存放配置。关键字段："
)

env_items = [
    "DEEPSEEK_API_KEY=sk-xxx    # DeepSeek API Key",
    "DEEPSEEK_MODEL=deepseek-chat  # 模型名称",
    "DEEPSEEK_BASE_URL=https://api.deepseek.com  # API 地址",
]
for c in env_items:
    p = doc.add_paragraph(c)
    for run in p.runs:
        run.font.name = "Consolas"
        run.font.size = Pt(10)

doc.add_heading("4.4 talk 命令详解", level=2)
doc.add_paragraph(
    "talk 是聊天模式命令，专为日常对话场景设计。与 shell 模式不同，talk 模式："
)
talk_features = [
    "完全隐藏工具调用过程：后台自动搜索/执行代码/读写文件，用户看不到任何技术细节",
    "拟人化回答风格：有情感、有温度、像真人朋友一样聊天",
    "时效性优先：问起当前事件、体育比赛、天气、新闻、赛程等，必须在后台先联网搜索，再回答",
    "主动互动：回答后偶尔自然关心用户想法，抛出延伸问题",
    "默认温度 1.0：回答更活泼、有创意",
]
for f in talk_features:
    doc.add_paragraph(f, style="List Bullet")
doc.add_paragraph(
    "talk 模式需要 DEEPSEEK_API_KEY。启动方式：polyagent talk。"
    "输入 exit 或 Ctrl-D 退出。"
)

doc.add_page_break()

# ══════════════════════════════════════════════════════
# 5. SDK 使用指南
# ══════════════════════════════════════════════════════
doc.add_heading("5. SDK 使用指南", level=1)

doc.add_heading("5.1 单 Agent — 基础对话", level=2)

code1 = '''from polyagent.core import Agent, AgentSpec
from polyagent.llm import LLMClient, DeepSeekProvider

agent = Agent(
    AgentSpec(name="assistant", role="assistant", model="deepseek-chat"),
    LLMClient(DeepSeekProvider()),
)
resp = await agent.run("你好，请介绍一下自己")
print(resp.content)'''
p = doc.add_paragraph(code1)
for run in p.runs:
    run.font.name = "Consolas"
    run.font.size = Pt(9)

doc.add_heading("5.2 单 Agent + 工具", level=2)

code2 = '''from polyagent.tools import with_builtins

agent = Agent(
    AgentSpec(name="coder", role="worker"),
    LLMClient(DeepSeekProvider()),
    tools=with_builtins(),  # 挂载全部工具
)
resp = await agent.run("计算 1+2+...+100 并保存到 result.txt")
print(resp.content)'''
p = doc.add_paragraph(code2)
for run in p.runs:
    run.font.name = "Consolas"
    run.font.size = Pt(9)

doc.add_heading("5.3 多 Agent 编排", level=2)

code3 = '''from polyagent.orchestration import (
    Orchestrator, Planner, Worker, Critic, Synthesizer
)
from polyagent.observability import Tracer
from polyagent.core.types import CostReport

planner = Planner(agent_planner)
worker = Worker(agent_worker)
critic = Critic(agent_critic)
synthesizer = Synthesizer(agent_synthesizer)

orchestrator = Orchestrator(
    planner, worker, critic, synthesizer,
    tracer=Tracer(), cost_report=CostReport(),
    max_review_retries=2,
)
result = await orchestrator.run("分析当前项目的代码质量")
print(f"答案: {result.answer}")
print(f"成本: ${result.estimated_cost_usd:.6f}")
print(f"耗时: {result.latency:.3f}s")'''
p = doc.add_paragraph(code3)
for run in p.runs:
    run.font.name = "Consolas"
    run.font.size = Pt(9)

doc.add_heading("5.4 自定义工具", level=2)

code4 = '''from polyagent.tools import Tool
from pydantic import BaseModel

class MyArgs(BaseModel):
    url: str

class MyTool(Tool):
    name = "my_tool"
    description = "做某件事"
    args_model = MyArgs

    async def run(self, args: MyArgs) -> ToolResult:
        return ToolResult(output=f"done: {args.url}")

# 注册到 ToolRegistry
from polyagent.tools import ToolRegistry
reg = ToolRegistry()
reg.register(MyTool())'''
p = doc.add_paragraph(code4)
for run in p.runs:
    run.font.name = "Consolas"
    run.font.size = Pt(9)

doc.add_page_break()

# ══════════════════════════════════════════════════════
# 6. 内置工具一览
# ══════════════════════════════════════════════════════
doc.add_heading("6. 内置工具一览", level=1)

tools_info = [
    ("python_execute", "执行 Python 代码，返回 stdout/stderr", "计算、数据处理、系统信息"),
    ("read_file", "读取文件内容（UTF-8 文本）", "查看代码、配置文件"),
    ("write_file", "将文本写入文件（覆写或追加）", "保存内容、创建文件"),
    ("http_request", "发送 HTTP 请求", "调 API、抓取网页、下载内容"),
    ("grep_files", "递归搜索文件中的正则模式", "代码搜索、日志分析"),
    ("web_search", "通过 DuckDuckGo 搜索网络", "实时信息、新闻、文档查找"),
    ("marketplace_search", "搜索内建技能仓库", "查找可安装的 PolyAgent 技能"),
    ("marketplace_install", "安装内建技能", "安装天气/文件分析/时间工具"),
    ("skillhub_search", "搜索腾讯 SkillHub 市场（76K+ 技能）", "查找 PDF/代码/数据处理等技能"),
    ("skillhub_install", "按 slug 安装 SkillHub 技能", "下载技能并解压到本地"),
    ("skillhub_install_from_prompt", "从官方安装 prompt 安装技能", "粘贴官网 prompt 自动安装"),
    ("skillhub_list", "列出已安装的 SkillHub 技能", "查看本地安装的技能"),
]

table2 = doc.add_table(rows=1, cols=3)
table2.style = "Light Grid Accent 1"
table2.alignment = WD_TABLE_ALIGNMENT.CENTER

hdr2 = table2.rows[0].cells
hdr2[0].text = "工具名"
hdr2[1].text = "功能说明"
hdr2[2].text = "适用场景"

for name, desc, scene in tools_info:
    row = table2.add_row().cells
    row[0].text = name
    row[1].text = desc
    row[2].text = scene

doc.add_paragraph()

doc.add_page_break()

# ══════════════════════════════════════════════════════
# 7. 技能市场
# ══════════════════════════════════════════════════════
doc.add_heading("7. 技能市场", level=1)

doc.add_heading("7.1 PolyAgent 内建市场", level=2)
doc.add_paragraph(
    "PolyAgent 自带 3 个内建技能，使用 marketplace_search / marketplace_install 操作："
)
builtins = [
    ("weather", "get_weather(city)", "查询任意城市的天气（模拟数据）"),
    ("file_analyzer", "analyze_file(path)", "分析文件行数/字数/大小/编码"),
    ("datetime_utils", "current_time() / date_diff() / format_timestamp()", "日期时间工具"),
]
for name, tools, desc in builtins:
    doc.add_paragraph(f"• {name}：{tools} — {desc}")

doc.add_heading("7.2 腾讯 SkillHub 市场", level=2)
doc.add_paragraph(
    "腾讯 SkillHub（https://skillhub.tencent.com）是 WorkBuddy 生态的官方技能商店，"
    "拥有 76,000+ 个 AI 技能。PolyAgent 通过 REST API 直接访问，无需安装 CLI。"
)
doc.add_paragraph("使用方式：")
skillhub_ways = [
    "skillhub_search('关键词') — 搜索技能",
    "skillhub_install('slug') — 按标识符安装",
    "skillhub_install_from_prompt('粘贴的prompt') — 从官网 prompt 安装",
    "skillhub_list() — 查看已安装的技能",
]
for w in skillhub_ways:
    doc.add_paragraph(w, style="List Bullet")

doc.add_paragraph(
    "安装的技能以 SKILL.md 文件形式存放在 ~/.polyagent/skillhub-skills/{slug}/ 目录下。"
    "Agent 可通过 read_file 读取说明了解技能用法。"
)

doc.add_page_break()

# ══════════════════════════════════════════════════════
# 8. 多 Agent 编排
# ══════════════════════════════════════════════════════
doc.add_heading("8. 多 Agent 编排详解", level=1)

doc.add_heading("8.1 角色定义", level=2)
roles = [
    ("Planner", "将用户目标分解为 DAG（有向无环图）子任务。每个任务有 id、description、deps（依赖列表）。"),
    ("Worker", "并发执行分配的子任务。可挂载工具（python_execute、read_file 等），支持自动重试。"),
    ("Critic", "评审 Worker 的输出。输出 JSON：{\"accepted\": true/false, \"feedback\": \"...\"}。不通过回退重做。"),
    ("Synthesizer", "汇总所有任务结果，生成最终答案。"),
]
for role, desc in roles:
    p = doc.add_paragraph()
    run = p.add_run(f"• {role}：")
    run.bold = True
    p.add_run(desc)

doc.add_heading("8.2 任务调度", level=2)
doc.add_paragraph(
    "Orchestrator 使用拓扑排序调度任务。无依赖的任务并行执行，依赖链按序执行。"
    "每个 Worker 任务配置 Semaphore 控制并发度（默认 5）。"
    "Critic 评审不通过时，Worker 重做，最多重试 max_review_retries 次。"
)

doc.add_heading("8.3 错误处理", level=2)
doc.add_paragraph(
    "编排层内置错误处理机制："
    "Worker 失败（返回 error=True）的任务进入 failed 状态，"
    "其下游任务（直接或间接依赖它的）自动标记为 blocked 不会被调度。"
    "Planner 返回无效 JSON 格式时，编排器优雅跳过。"
)

doc.add_page_break()

# ══════════════════════════════════════════════════════
# 9. 企业级特性
# ══════════════════════════════════════════════════════
doc.add_heading("9. 企业级特性", level=1)

doc.add_heading("9.1 中间件链", level=2)
doc.add_paragraph(
    "LLMClient 支持链式中间件。所有 LLM 请求依次经过："
)
middleware = [
    "限流（Throttle）：控制请求速率",
    "重试（Retry）：失败自动重试，指数退避",
    "降级（Fallback）：Provider 不可用时切换到备用",
    "预算校验（Budget）：超过预算限制请求",
    "成本统计（CostAccount）：累加 token 使用和费用",
]
for m in middleware:
    doc.add_paragraph(m, style="List Bullet")

doc.add_heading("9.2 RAG", level=2)
doc.add_paragraph(
    "RAG（检索增强生成）子系统包含："
    "FastEmbedEmbedder（paraphrase-multilingual-MiniLM ONNX 本地语义模型，无需联网）"
    "和 ChromaVectorStore（ChromaDB 持久化向量存储，SQLite 后端）。"
    "支持文本分块（RecursiveCharacterSplitter）、索引管理。"
    "不使用 DeepSeek 的 embedding API（实测 DeepSeek 不支持 embedding 接口）。"
)

doc.add_heading("9.3 可观测性", level=2)
doc.add_paragraph(
    "可观测性基于 contextvars 实现无侵入 Tracing。"
    "Span 树自动记录请求链路（父-子关系），支持导出到 Jaeger（OTLP 协议）"
    "和 Prometheus（pushgateway 模式）。"
    "内置 docker-compose 配置一键启动 Jaeger + pushgateway + Prometheus。"
    "已通过真实 Jaeger 联调验证（10+ spans，含 planner.plan / schedule / synthesizer.synthesize 等）。"
)

doc.add_heading("9.4 Docker 沙箱", level=2)
doc.add_paragraph(
    "tools/sandbox.py 实现 DockerSandbox：使用 docker run --rm --network=none "
    "--memory=256m --cpus=0.5 python:3.12-slim 容器隔离执行不可信代码。"
    "容器不可访问网络和主机文件系统。"
    "PathGuard 提供路径白名单防护。"
    "已通过真实 Docker 联调验证。"
)

doc.add_heading("9.5 评估框架", level=2)
doc.add_paragraph(
    "Eval 子系统支持自动化质量评估："
    "ContainsScorer（检查输出是否包含关键字）、"
    "ExactMatchScorer（精确匹配）、"
    "LLMJudgeScorer（调用 DeepSeek 作为评分裁判，返回 JSON {score, reason}）。"
    "EvalRunner 通过 subject 接口与任意 Agent 解耦。"
    "命令行 polyagent eval 执行内置数据集。"
)

doc.add_heading("9.6 持久化", level=2)
doc.add_paragraph(
    "persistence/store.py 基于 SQLite 实现 RunStore。"
    "每次 polyagent run 自动保存目标、结果、任务图、Trace。"
    "支持 list（最近 20 条）和 get（按 ID 查详情）操作。"
    "CLI 命令 runs / show 提供查询接口。"
)

doc.add_page_break()

# ══════════════════════════════════════════════════════
# 10. 测试与质量保证
# ══════════════════════════════════════════════════════
doc.add_heading("10. 测试与质量保证", level=1)

doc.add_paragraph("当前测试状态（89 个单测通过）：")

qa_table = doc.add_table(rows=1, cols=3)
qa_table.style = "Light Grid Accent 1"
qa_table.alignment = WD_TABLE_ALIGNMENT.CENTER

hdr3 = qa_table.rows[0].cells
hdr3[0].text = "指标"
hdr3[1].text = "当前值"
hdr3[2].text = "门禁"

qa_data = [
    ("测试通过数", "89", "≥ 55"),
    ("代码覆盖率", "75%", "≥ 75%"),
    ("mypy strict", "0 错误", "0"),
    ("ruff", "全绿", "全绿"),
    ("类型标注", "全员", "全员"),
]
for name, val, threshold in qa_data:
    row = qa_table.add_row().cells
    row[0].text = name
    row[1].text = val
    row[2].text = threshold

doc.add_paragraph()

doc.add_paragraph("CI 流程包含 mypy strict 检查、pytest-cov 覆盖率门禁（75%）、ruff lint 检查。")
doc.add_paragraph("Docker 交付：基于 python:3.12-slim 的多阶段构建镜像，通过 docker-compose 一键部署。")

doc.add_page_break()

# ══════════════════════════════════════════════════════
# 11. 常见问题
# ══════════════════════════════════════════════════════
doc.add_heading("11. 常见问题", level=1)

faqs = [
    ("Q: 为什么 polyagent 命令找不到？",
     "A: polyagent.exe 安装在 venv 的 Scripts 目录下，不在系统 PATH 中。"
     "解决方案：set PATH=<venv>/Scripts;%PATH% 或使用全路径。"),
    ("Q: 为什么 Agent 的回答只是重复我的输入？",
     "A: 默认使用 MockProvider（离线模拟），它原样返回输入。"
     "加 --provider deepseek --stream 使用真实 LLM。"),
    ("Q: 为什么 chat 感觉和普通 LLM 一样？",
     "A: chat 命令没有挂载工具。用 shell 命令可以获得带全部工具的智能助手。"),
    ("Q: shell 和 talk 有什么区别？",
     "A: shell 是技术模式，显示工具调用过程和结构化回答，适合开发/调试场景。"
     "talk 是聊天模式，隐藏所有后台操作，以拟人化风格回答，适合日常对话。"
     "两者使用相同的工具集，只是呈现方式不同。"),
    ("Q: .env 文件放哪？应该在哪里运行 polyagent 命令？",
     "A: 两个关键点：(1) .env 文件必须放在项目根目录（polyagent 在当前工作目录下查找 .env）。"
     "(2) polyagent.exe 安装在 venv 的 Scripts 目录下，不在系统 PATH 中。"
     "所以正确的用法是：先 cd 到项目根目录（D:\\孙\\WorkBuddy\\2026-07-10-18-40-02），"
     "再用 set PATH=... 将 Scripts 加入 PATH，最后执行 polyagent 命令。"
     "详细步骤见第 3 章「快速开始」的 3.2 节。"),
    ("Q: 如何给 Agent 添加自定义工具？",
     "A: 继承 Tool 基类，实现 run 方法，然后注册到 ToolRegistry。参考 SDK 文档。"),
    ("Q: SkillHub 技能装到哪了？",
     "A: ~/.polyagent/skillhub-skills/{slug}/ 目录下，包含 SKILL.md 文件。"),
    ("Q: 支持哪些 LLM 提供商？",
     "A: 内置 MockProvider（离线）和 DeepSeekProvider。任何 OpenAI 兼容协议均可通过继承 BaseProvider 接入。"),
]
for q, a in faqs:
    p = doc.add_paragraph()
    run = p.add_run(q)
    run.bold = True
    doc.add_paragraph(a)
    doc.add_paragraph()

doc.add_page_break()

# ══════════════════════════════════════════════════════
# 12. 附录
# ══════════════════════════════════════════════════════
doc.add_heading("12. 附录", level=1)

doc.add_heading("12.1 项目结构", level=2)
tree = """polyagent/
├── __init__.py         # 版本号
├── cli/                # CLI 命令实现
│   ├── __init__.py     # version/run/chat/shell/eval/runs/show
│   └── demo.py         # 演示编排器
├── core/               # 核心类型和 Agent
│   ├── agent.py        # Agent（单 Agent + 工具循环）
│   ├── types.py        # Message/Role/ToolCall/Usage/AgentSpec
│   └── exceptions.py   # LLMError/RateLimitError/TimeoutError
├── llm/                # LLM 提供商
│   ├── client.py       # LLMClient（中间件链）
│   ├── mock.py         # MockProvider（离线）
│   ├── deepseek.py     # DeepSeekProvider（真实）
│   └── middleware.py   # 限流/重试/降级/预算/成本
├── tools/              # 工具系统
│   ├── base.py         # Tool 基类
│   ├── builtins.py     # 12 个内置工具
│   └── sandbox.py      # DockerSandbox + PathGuard
├── skills/             # 技能市场
│   ├── registry.py     # 内建注册表
│   ├── installer.py    # 安装/卸载/加载
│   └── builtins/       # 示例技能
├── orchestration/      # 编排层
│   ├── orchestrator.py # Orchestrator
│   ├── roles.py        # Planner/Worker/Critic/Synthesizer
│   └── types.py        # TaskNode/Critique/RunResult
├── observability/      # 可观测
│   ├── tracer.py       # Span 树
│   └── metrics.py      # 指标采集
├── memory/             # 记忆系统
├── rag/                # RAG 子系统
├── eval/               # 评估框架
├── persistence/        # SQLite 持久化
└── config.py           # 配置读取"""
p = doc.add_paragraph(tree)
for run in p.runs:
    run.font.name = "Consolas"
    run.font.size = Pt(8)

doc.add_paragraph()

doc.add_heading("12.2 技术栈", level=2)
stack = [
    ("运行环境", "Python ≥ 3.11, asyncio"),
    ("HTTP 客户端", "httpx (async)"),
    ("数据校验", "pydantic"),
    ("CLI 框架", "typer"),
    ("向量检索", "fastembed (ONNX), chromadb"),
    ("可观测", "opentelemetry-sdk, prometheus-client"),
    ("文档生成", "python-docx"),
]
table3 = doc.add_table(rows=1, cols=2)
table3.style = "Light Grid Accent 1"
table3.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr4 = table3.rows[0].cells
hdr4[0].text = "类目"
hdr4[1].text = "技术"

for cat, tech in stack:
    row = table3.add_row().cells
    row[0].text = cat
    row[1].text = tech

doc.add_paragraph()

doc.add_heading("12.3 许可", level=2)
doc.add_paragraph(
    "PolyAgent 是一个开源项目，遵循本仓库的许可证条款。"
    "腾讯 SkillHub 的技能版权归其各自作者所有。"
)

# ── 保存 ──────────────────────────────────────────────
output = "D:\\孙\\WorkBuddy\\2026-07-10-18-40-02\\PolyAgent_产品说明书.docx"
try:
    doc.save(output)
    print(f"OK: saved to {output}")
except PermissionError:
    # If file is locked (e.g. opened in Word/WPS), save with a timestamp suffix
    from datetime import datetime
    alt = output.replace(".docx", f"_{datetime.now().strftime('%H%M%S')}.docx")
    doc.save(alt)
    print(f"OK: saved to {alt} (original was locked)")
